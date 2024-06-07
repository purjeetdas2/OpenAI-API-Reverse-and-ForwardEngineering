
from docx import Document
from docx.shared import Pt


class WordDocumentGenerator:
    def __init__(self, json_data, file_name='Output.docx'):
        self.json_data = json_data
        self.file_name = file_name
        self.document = Document()

    def _parse_dict(self, d, indent_level, top_level=False):
        if isinstance(d, dict):
            for key, value in d.items():
                if top_level:
                    self.document.add_heading(key, level=1)
                    self._parse_value(value, indent_level + 1)
                else:
                    para = self.document.add_paragraph()
                    run = para.add_run(f'{key}: ')
                    run.bold = True
                    if not isinstance(value, dict) and not isinstance(value, list):
                      para.add_run(str(value))
                      para.paragraph_format.left_indent = Pt(18) * indent_level
                      para.paragraph_format.space_before = Pt(0)
                      para.paragraph_format.space_after = Pt(0)
                    else:
                      para.paragraph_format.left_indent = Pt(18) * indent_level
                      para.paragraph_format.space_before = Pt(0)
                      para.paragraph_format.space_after = Pt(0)
                      self._parse_value(value, indent_level + 1)                   
                    #self._parse_value(value, indent_level + 1)
        else:
            self._parse_value(d, indent_level + 1)

    def _parse_value(self, value, indent_level):
        if isinstance(value, dict):
            self._parse_dict(value, indent_level)
        elif isinstance(value, list):
            self._parse_list(value, indent_level)
            
        else:
            para = self.document.add_paragraph(str(value))
            para.paragraph_format.left_indent = Pt(18) * indent_level
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)

    def _parse_list(self, l, indent_level):
        for item in l:
            if isinstance(item, dict):
                self._parse_dict(item, indent_level)
            else:
                para = self.document.add_paragraph(f'- {item}')
                para.paragraph_format.left_indent = Pt(18) * indent_level
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
    def generate_document(self):
        """
        Generate the Word document
        """
        # Process only the top-level dictionary with headings
         # Start parsing top-level dictionary with base indentation
        self._parse_dict(self.json_data, indent_level=0, top_level=True)
        self.document.save(self.file_name)