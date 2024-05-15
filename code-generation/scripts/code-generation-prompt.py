from openai import OpenAI
import docx
import os


client = OpenAI()
prompt = """
You are an expert software developer proficient in both mainframe technologies and modern Java-based frameworks, particularly Spring Batch. Given a document that summarizes the key elements of a mainframe system, including JCL Jobs, COBOL programs, Subroutines, and Copybooks, your task is to convert these business rules and logic into a Spring Batch application. 

Content of the document:
- High-level summary of JCL Job functionalities
- Details of COBOL program logic
- Description of Subroutine workflows
- Content and structure of Copybooks
"""
def extract_code(bre):
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
        {
            "role": "user",
            "content": prompt+bre
        }
        ],
        temperature=0.7,
        max_tokens=4000,
        top_p=0.95,
        frequency_penalty=0,
        presence_penalty=0
    )
    return response.choices[0].message.content
def extract_bre(input_folder):
    print(f"..... Read the code from the input folder location ...[{input_folder}]")
    for filename in os.listdir(input_folder):
        file_path= os.path.join(input_folder,filename)
        if os.path.isfile(file_path) and file_path.endswith(".docx"):
            print(file_path)
            doc = docx.Document(file_path)
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
    return '\n'.join(fullText)

def save_to_code_location(content,target_file_name):
    print(f"..... Writing the extracted code to file - [{target_file_name}]")
    doc = docx.Document()
    doc.add_paragraph(content)
    doc.save(target_file_name)


def main():
    current_working_directory = os.getcwd()
    try:
        bre_doc = extract_bre(input_folder=current_working_directory+'/docs/')
        extracted_code = extract_code(bre=bre_doc)
        save_to_code_location(content=extracted_code,target_file_name=current_working_directory+'/code/'+'SampleCodeGenrated.docx')
        print("...... Code Generated successfully ...")
    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    main()