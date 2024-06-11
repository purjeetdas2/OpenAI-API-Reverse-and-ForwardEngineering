import requests
import openai
import os
import docx
from openai import OpenAI
import re

class CodeGenerator:
    def __init__(self, api_key: str, prompt: str, service: str = 'openai', endpoint: str = None, organization: str = None):
        self.service = service.lower()
        self.prompt = prompt
        self.api_key = api_key
        self.endpoint = endpoint

        if self.service == 'openai':
            self.client = OpenAI(api_key=api_key,organization=organization)
        elif self.service == 'azure' and not endpoint:
            raise ValueError("Azure endpoint must be provided for Azure OpenAI service")

    def _create_chat_messages(self, bre: str) -> list:
        return [
            {
                "role": "system",
                "content": "You are an expert software developer proficient in both mainframe technologies and modern Java-based frameworks, particularly Spring Batch. Your task is to convert given business rules and logic from mainframe systems into a Spring Batch application."
            },
            {
                "role": "user",
                "content": self.prompt + bre
            },
            {
                "role": "assistant",
                "content": "Please provide: \n1. Translation of JCL Job functionalities to Spring Batch configurations. \n2. Java equivalents of COBOL program, Subroutine logic using Spring Batch framework. \n3. Mapping of Subroutine workflows to Spring Batch components. \n4. Conversion of Copybooks content and structure to equivalent Java classes or data structures used in Spring Batch. Follow the best practices and naming conventions in Spring batch. Generate unit test classes required to test the spring batch."
            }
        ]

    def extract_code(self, bre: str) -> str:
        messages = self._create_chat_messages(bre)
        response = None

        if self.service == 'azure':
            headers = {
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {self.api_key}'
            }
            data = {
                "model": "gpt-4-omni",
                "messages": messages,
                "temperature": 0.7,
                "max_tokens": 4000,
                "top_p": 0.95,
                "frequency_penalty": 0,
                "presence_penalty": 0
            }
            response = requests.post(f"{self.endpoint}/openai/deployments/gpt-4-omni/completions", headers=headers, json=data)
            response.raise_for_status()  # Raise an error for bad status codes
            result = response.json()
            return result['choices'][0]['message']['content']
        
        elif self.service == 'openai':
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=messages,
                temperature=0.7,
                max_tokens=4000,
                top_p=0.95,
                frequency_penalty=0,
                presence_penalty=0
            )
            return response.choices[0].message.content

    @staticmethod
    def extract_bre(input_folder: str) -> str:
        print(f"Read the bre from the input folder location ...[{input_folder}]")
        for filename in os.listdir(input_folder):
            file_path = os.path.join(input_folder, filename)
            if os.path.isfile(file_path) and file_path.endswith(".docx"):
                print(file_path)
                doc = docx.Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                return '\n'.join(full_text)
        raise FileNotFoundError(f"No .docx files found in the provided directory: {input_folder}")

    def save_to_code_location(self, content: str, target_folder: str,target_file_name: str):
        
        content="""
        Writing the extracted code to file - [/Users/purjeetdas/development/GenAI-repo/OpenAI-AzureAI/OpenAI-API-Reverse-and-ForwardEngineering/code-generation/code/edi-850-vendor-XEDI850_20240610_161235.docx]
        ### 1. Spring Batch Configuration (Translation of JCL Job Functionalities)

        **High-Level Code Flow Outline**

        ```java
        @Configuration
        @EnableBatchProcessing
        public class BatchConfig {
            
            @Autowired
            private JobBuilderFactory jobBuilderFactory;
            
            @Autowired
            private StepBuilderFactory stepBuilderFactory;
            
            @Autowired
            private DataSource dataSource;
            
            @Bean
            public FlatFileItemReader<InputRecord> reader() {
                return new FlatFileItemReaderBuilder<InputRecord>()
                    .name("inputRecordReader")
                    .resource(new FileSystemResource("/path/to/input/file.txt"))
                    .delimited()
                    .names(new String[]{"fieldA", "fieldB", "fieldC", "fieldD", "fieldE", "fieldF"})
                    .targetType(InputRecord.class)
                    .build();
            }
            
            @Bean
            public ItemProcessor<InputRecord, OutputRecord> processor() {
                return new BusinessRulesProcessor();
            }
            
            @Bean
            public FlatFileItemWriter<OutputRecord> writer() {
                return new FlatFileItemWriterBuilder<OutputRecord>()
                    .name("outputRecordWriter")
                    .resource(new FileSystemResource("/path/to/output/file.txt"))
                    .delimited()
                    .names(new String[]{"fieldA", "fieldB", "fieldC", "fieldD", "fieldE", "fieldF"})
                    .build();
            }
            
            @Bean
            public Job importUserJob(JobCompletionNotificationListener listener, Step step1) {
                return jobBuilderFactory.get("importUserJob")
                    .incrementer(new RunIdIncrementer())
                    .listener(listener)
                    .flow(step1)
                    .end()
                    .build();
            }
            
            @Bean
            public Step step1(FlatFileItemWriter<OutputRecord> writer) {
                return stepBuilderFactory.get("step1")
                    .<InputRecord, OutputRecord>chunk(10)
                    .reader(reader())
                    .processor(processor())
                    .writer(writer)
                    .build();
            }
        }
        ```

        ### 2. Business Validation Rules (Java Equivalent of COBOL Program Logic)

        ```java
        public class BusinessRulesProcessor implements ItemProcessor<InputRecord, OutputRecord> {
            
            @Override
            public OutputRecord process(final InputRecord input) throws Exception {
                validate(input);
                return applyBusinessRules(input);
            }

            private void validate(InputRecord input) {
                // Rule 1: Field 'A' must be numeric and non-null
                if (input.getFieldA() == null || input.getFieldA() < 0) {
                    throw new ValidationException("Field 'A' is invalid");
                }

                // Rule 2: Field 'B' must not exceed 10 characters
                if (input.getFieldB().length() > 10) {
                    throw new ValidationException("Field 'B' exceeds maximum length");
                }
            }
            
            private OutputRecord applyBusinessRules(InputRecord input) {
                OutputRecord output = new OutputRecord();

                // Copy input to output
                output.setFieldA(input.getFieldA());
                output.setFieldB(input.getFieldB());
                output.setFieldC(input.getFieldC());
                output.setFieldE(input.getFieldE());
                output.setFieldF(input.getFieldF());

                // Apply Rule 1: If Field 'C' is 'Y', set Field 'D' to current date
                if ("Y".equals(input.getFieldC())) {
                    output.setFieldD(new SimpleDateFormat("yyyy-MM-dd").format(new Date()));
                }

                // Apply Rule 2: Multiply Field 'E' by 1.2 if Field 'F' is greater than 100
                if (input.getFieldF() > 100) {
                    output.setFieldE(input.getFieldE() * 1.2);
                }

                return output;
            }
        }
        ```

        ### 3. Copybook Data Structures (Java Classes)

        ```java
        public class InputRecord {
            private Integer fieldA;
            private String fieldB;
            private String fieldC;
            private String fieldD;
            private Integer fieldE;
            private Integer fieldF;

            // Getters and Setters
            public Integer getFieldA() { return fieldA; }
            public void setFieldA(Integer fieldA) { this.fieldA = fieldA; }

            public String getFieldB() { return fieldB; }
            public void setFieldB(String fieldB) { this.fieldB = fieldB; }

            public String getFieldC() { return fieldC; }
            public void setFieldC(String fieldC) { this.fieldC = fieldC; }

            public String getFieldD() { return fieldD; }
            public void setFieldD(String fieldD) { this.fieldD = fieldD; }

            public Integer getFieldE() { return fieldE; }
            public void setFieldE(Integer fieldE) { this.fieldE = fieldE; }

            public Integer getFieldF() { return fieldF; }
            public void setFieldF(Integer fieldF) { this.fieldF = fieldF; }
        }

        public class OutputRecord {
            private Integer fieldA;
            private String fieldB;
            private String fieldC;
            private String fieldD;
            private Double fieldE;
            private Integer fieldF;

            // Getters and Setters
            public Integer getFieldA() { return fieldA; }
            public void setFieldA(Integer fieldA) { this.fieldA = fieldA; }

            public String getFieldB() { return fieldB; }
            public void setFieldB(String fieldB) { this.fieldB = fieldB; }

            public String getFieldC() { return fieldC; }
            public void setFieldC(String fieldC) { this.fieldC = fieldC; }

            public String getFieldD() { return fieldD; }
            public void setFieldD(String fieldD) { this.fieldD = fieldD; }

            public Double getFieldE() { return fieldE; }
            public void setFieldE(Double fieldE) { this.fieldE = fieldE; }

            public Integer getFieldF() { return fieldF; }
            public void setFieldF(Integer fieldF) { this.fieldF = fieldF; }
        }
        ```

        ### 4. Control Card Information (Used in Spring Batch Configuration)

        The control card information is utilized in the `FlatFileItemReader` and `FlatFileItemWriter` bean configurations to set file paths and table names.

        ### 5. Unit Test Classes

        ```java
        @RunWith(SpringRunner.class)
        @SpringBootTest
        public class BatchTest {
            
            @Autowired
            private JobLauncherTestUtils jobLauncherTestUtils;
            
            @Test
            public void testJob() throws Exception {
                JobExecution jobExecution = jobLauncherTestUtils.launchJob();
                assertEquals(BatchStatus.COMPLETED, jobExecution.getStatus());
            }
            
            @Test
            public void testStep() {
                JobExecution jobExecution = jobLauncherTestUtils.launchStep("step1");
                assertEquals(BatchStatus.COMPLETED, jobExecution.getStatus());
            }
            
            @Test
            public void testProcessor() throws Exception {
                BusinessRulesProcessor processor = new BusinessRulesProcessor();
                InputRecord input = new InputRecord();
                input.setFieldA(12345);
                input.setFieldB("TestField");
                input.setFieldC("Y");
                input.setFieldD("2023-10-01");
                input.setFieldE(120);
                input.setFieldF(150);
                
                OutputRecord output = processor.process(input);
                assertNotNull(output);
                assertEquals("2023-10-01", output.getFieldD());
                assertEquals(144.0, output.getFieldE(), 0.01);
            }
        }
        ```

        ### Generated Filenames

        ```
        generated files: BatchConfig.java, BusinessRulesProcessor.java, InputRecord.java, OutputRecord.java, BatchTest.java
        ```

         """
        print(f"Writing the extracted code to file - [{target_file_name}]")
        doc = docx.Document()
        print(content)
        doc.add_heading('Response:', level=1)

        filePattern = re.compile(r'```generated files(.*?)```', re.DOTALL)
        files_generated = filePattern.findall(content)
        print(files_generated)
        # Directory to store the generated Java files
        # Pattern to extract Java code blocks based on typical Java code markers
        # We use re.DOTALL to allow the dot toß match newlines

        # Define the output directory for generated Java files
        """
        output_dir = os.path.join(target_folder,"generated_java_files")
        os.makedirs(output_dir, exist_ok=True)
        pattern = re.compile(r'```java(.*?)```', re.DOTALL)

        # Extract all Java code blocks
        code_blocks = pattern.findall(content)

        
        print(code_blocks)
        # Extracting content for each file based on the order and known structure of the response
        batch_config_content = code_blocks[0].strip()
        business_rules_processor_content = code_blocks[1].strip()
        input_record_content = code_blocks[2].strip()
        #output_record_content = code_blocks[3].strip()
        unit_tests_content = code_blocks[3].strip()

        # File names and contents mapping
        files = {
            "BatchConfig.java": batch_config_content,
            "BusinessRulesProcessor.java": business_rules_processor_content,
            "InputRecord.java": input_record_content,
            "BatchConfigTest.java": unit_tests_content,
        }

        # Write to files
        for file_name, content in files.items():
            with open(os.path.join(output_dir, file_name), "w") as file:
                file.write(content)

        print(f"Java files have been generated in {os.path.abspath(output_dir)}")

        paragraphs = content.split('\n')
        for paragraph in paragraphs:
            clean_paragraph = paragraph.strip()
            if clean_paragraph.startswith('###'):
                doc.add_heading(clean_paragraph.strip('# ').title(), level=2)
            elif clean_paragraph.startswith('####'):
                doc.add_heading(clean_paragraph.strip('# ').title(), level=3)
            elif clean_paragraph.startswith('-'):
                p = doc.add_paragraph(clean_paragraph.lstrip('- '))
                p.style = 'List Bullet'
            elif clean_paragraph:
                self._add_paragraph_with_formatting(doc, clean_paragraph)
        
       
        print(f"Code Generation is complete and  files have been generated in {os.path.abspath(output_dir)}")
        doc.save(target_file_name)
        """

    @staticmethod
    def _add_paragraph_with_formatting(doc: docx.Document, text: str):
        paragraph = doc.add_paragraph()
        parts = text.split('**')
        bold = False
        for part in parts:
            run = paragraph.add_run(part)
            if bold:
                run.bold = True
            bold = not bold
    